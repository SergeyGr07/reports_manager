from flask import Blueprint, render_template, request, jsonify
import os
from config import add_logger
from flask import send_file
from docx import Document
from io import BytesIO
import numpy as np
# from docx.shared import Inches
# from openpyxl import load_workbook
import json


script_name = os.path.splitext(os.path.basename(__file__))[0]
logger = add_logger(f'logger_{script_name}', script_name)


report = Blueprint("reports", __name__)


AVERAGE_COLUMN = 'Average'
ALLOWANCE_COLUMN = 'Allowance'
GEOMETRY_COLUMN = 'Geometry'


@report.route('/ticket_form')
def ticket_form():
    return render_template('ticket.html')


@report.route('/select_detail', methods=['POST', 'GET'])
def select_detail():

    details_map = {
        'Труба': {'col': 3, 'row': 11},
        'Кронштейн': {'col': 3, 'row': 6},
        'Крышка': {'col': 3, 'row': 16},
        'Тройник': {'col': 3, 'row': 7}
    }

    detail = request.json.get('detail')
    print(detail)

    response = details_map.get(detail, {})
    return jsonify(response), 200


# def add_columns_and_calculate_average(data_json):
#     # Функция для вычисления среднего значения списка чисел с помощью numpy
#     def calculate_average(numbers):
#         return np.mean(numbers) if numbers.size > 0 else None

#     # Добавление столбцов и вычисление среднего значения
#     for title, section in data_json.items():
#         for item in section.values():
#             header = item[0]
#             # Добавляем заголовки новых столбцов
#             header.extend(['Average', 'Allowance', 'Geometry'])
#             for row in item[1:]:
#                 measurements = np.array(row[2], dtype=np.float64)  # Преобразуем строки в числа с помощью numpy
#                 average = calculate_average(measurements)  # Вычисляем среднее значение с помощью numpy
#                 # Добавляем среднее значение и пустые значения для новых столбцов
#                 row.extend([average, '', ''])

#     return data_json


def create_table(data_dict):
    for section in data_dict.values():
        for item in section.values():
            header = item[0]
            if GEOMETRY_COLUMN not in header:
                header.extend([AVERAGE_COLUMN, ALLOWANCE_COLUMN, GEOMETRY_COLUMN])
            for row in item[1:]:
                measurements = np.array(row[2], dtype=np.float64)
                average = np.nanmean(measurements)
                row.extend([average, '', ''])
    return data_dict


@report.route('/put_data', methods=['POST'])
def save_data():
    try:
        data = request.get_json()
        print(data, '\n')
        if not data:
            return jsonify({'error': 'Отсутствуют данные в запросе'}), 400
        table_name = next(iter(data))

        print(create_table(data))

        with open(f'{table_name}.json', 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=4)
        return jsonify({'message': 'Данные успешно сохранены'}), 200

    except Exception as e:
        print(f"Ошибка при сохранении данных: {str(e)}")
        return jsonify({'error': 'Внутренняя ошибка сервера'}), 500


# @report.route('/submit_ticket', methods=['POST'])
# def submit_ticket():
#     title = request.form['ticket_title']
#     description = request.form['ticket_description']
#     num_rows = int(request.form['table_rows'])
#     num_cols = int(request.form['table_cols'])

#     return f'<h1>Ticket submitted:</h1><p>Title: {title}</p><p>Description: {description}</p><p>Table: {num_rows}x{num_cols}</p>'


@report.route('/')
def reports():
    return render_template('reports.html')


@report.route('/export_docx')
def fill_document():
    report_file = BytesIO()
    doc = Document("Template.docx")

    data = {
        'date': '21 марта 2024',
        'names': 'Налётов В.А, Петров П.П.',
        'name_machine': 'Оборудование X',
        'company': 'Компания ABC',
        'location': 'ул. Пушкина, д.10',
        'start_time': '10:00',
        'end_time': '12:00',
        'temp': '25',
    }

    data_pipe = {
        'Труба 1': {
            'nominal1': [356],
            'Measurements1': list(range(300, 311)),
            'Average1': [433],
            'Allowance1': ["что-то"],
            'Geometry1': ["что-то"]
        }
    }

    for paragraph in doc.paragraphs:
        for key, value in data.items():
            if '{{' + key + '}}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{{' + key + '}}', value)
        for key, value in data_pipe.items():
            if '{{' + key + '}}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{{' + key + '}}', value)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for name, pipe_data in data_pipe.items():
                    for key, value in pipe_data.items():
                        placeholder = '{{' + key + '}}'
                        if placeholder in cell.text:
                            cell.text = cell.text.replace(placeholder, ', '.join(map(str, value)))

    doc.save(report_file)
    report_file.seek(0)

    return send_file(report_file, as_attachment=True, download_name='report.docx', mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
